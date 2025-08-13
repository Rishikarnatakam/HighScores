import os
import json
import random
import re
from io import BytesIO
from dataclasses import dataclass
from typing import List, Dict, Any, Optional

import streamlit as st
from pydantic import BaseModel
from openai import OpenAI
from dotenv import load_dotenv


class Option(BaseModel):
    text: Optional[str] = ""
    isCorrect: bool = False


class GeneratedItem(BaseModel):
    question: Optional[str] = ""
    instruction: Optional[str] = ""
    difficulty: Optional[str] = "moderate"
    options: List[Option]
    explanation: Optional[str] = ""
    needsImage: Optional[bool] = False
    image_prompt: Optional[str] = ""


@dataclass
class TopicTriple:
    subject: str
    unit: str
    topic: str


def load_curriculum() -> List[TopicTriple]:
    with open(os.path.join('data', 'curriculum.json'), 'r', encoding='utf-8') as f:
        rows = json.load(f)
    return [TopicTriple(**r) for r in rows]


def _normalize_for_match(s: str) -> str:
    # Lowercase and collapse whitespace for rough duplicate detection
    return " ".join((s or "").lower().split())


def _fix_invalid_json_escapes(text: str) -> str:
    # Inside JSON strings, ensure backslashes start valid escapes; otherwise double them
    allowed = {'"', '\\', '/', 'b', 'f', 'n', 'r', 't', 'u'}
    out = []
    in_str = False
    escape = False
    quote = '"'
    i = 0
    while i < len(text):
        ch = text[i]
        if not in_str:
            if ch == '"':
                in_str = True
                out.append(ch)
            else:
                out.append(ch)
            i += 1
            continue
        # in string
        if escape:
            out.append(ch)
            escape = False
            i += 1
            continue
        if ch == '\\':
            nxt = text[i+1] if i + 1 < len(text) else ''
            if nxt in allowed:
                out.append(ch)
                escape = True
            else:
                out.append('\\\\')  # double the backslash
            i += 1
            continue
        # normalize raw control characters inside strings
        if ch == '\n' or ch == '\r':
            out.append('\\n')
            i += 1
            continue
        if ch == '\t':
            out.append('\\t')
            i += 1
            continue
        if ch == '"':
            in_str = False
            out.append(ch)
            i += 1
            continue
        out.append(ch)
        i += 1
    return ''.join(out)


def plan_distribution(triples: List[TopicTriple], total: int, seed: Optional[int]) -> List[Dict[str, Any]]:
    rng = random.Random(seed)
    shuffled = triples.copy()
    rng.shuffle(shuffled)
    T = len(shuffled)
    base = total // T
    remainder = total % T
    plan = [{"triple": t, "count": base} for t in shuffled]
    for i in range(remainder):
        plan[i % len(plan)]["count"] += 1
    return [p for p in plan if p["count"] > 0]


def read_first_env(names: List[str], default: str = "") -> str:
    for name in names:
        val = os.getenv(name)
        if val:
            return val
    return default


def build_client(base_url: str, api_key: str) -> OpenAI:
    return OpenAI(api_key=api_key, base_url=base_url)


DEFAULT_BASE_URL = "https://generativelanguage.googleapis.com/v1beta/openai/"


def call_model(client: OpenAI, model: str, system_prompt: str, user_prompt: str, temperature: float, reasoning_effort: str) -> str:
    # Single attempt only; no automatic fallbacks
    kwargs = dict(
        model=model,
        temperature=temperature,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    )
    # Gemini OpenAI-compatible API expects 'reasoning_effort': 'low' | 'medium' | 'high'
    if reasoning_effort and reasoning_effort.lower() != "none":
        kwargs["reasoning_effort"] = reasoning_effort.lower()
    resp = client.chat.completions.create(**kwargs)
    return resp.choices[0].message.content or ""


def build_system_prompt() -> str:
    return (
        "You are an expert math assessment author. "
        "Generate original multiple-choice questions that are curriculum-aligned and self-contained. "
        "Use LaTeX delimiters directly for math. Exactly 4 options with one correct. "
        "Provide concise correct explanations. "
        "Return only a JSON array and nothing else (no markdown, no prose). "
        "Use valid JSON with double-quoted keys/strings and no trailing commas. "
        "When using LaTeX inside JSON strings, escape backslashes (\\ becomes \\\\)."
    )
def _clean_field(text: str) -> str:
    if text is None:
        return ""
    t = str(text)
    # Fix escaped slashes common from markdown/JSON
    t = t.replace('\/', '/')
    # Normalize spaces around punctuation
    t = t.replace(' ,', ',').replace(' .', '.')
    t = re.sub(r'\s{2,}', ' ', t)
    return t.strip()


def _fix_inline_latex(text: str) -> str:
    if text is None:
        return ""
    t = str(text)
    # Convert \( ... \) and \[ ... \] to $...$ and $$...$$ for Streamlit markdown
    t = re.sub(r"\\\((.*?)\\\)", r"$\1$", t)
    t = re.sub(r"\\\[(.*?)\\\]", r"$$\1$$", t)
    # Wrap common LaTeX tokens not already inside $ ... $
    t = re.sub(r"(?<!\$)(\\frac\{[^}]+\}\{[^}]+\})(?!\$)", r"$\1$", t)
    t = re.sub(r"(?<!\$)(\\sqrt\{[^}]+\})(?!\$)", r"$\1$", t)
    t = re.sub(r"(?<!\$)([0-9]+)\^\{\\circ\}(?!\$)", r"$\1^{\\circ}$", t)
    # If we ended up with an odd number of $ delimiters, append one to balance
    if t.count('$') % 2 == 1:
        t = '$' + t
    return t



def build_user_prompt(triple: TopicTriple, count: int, style_context: str, allow_images: bool) -> str:
    ctx = ("\nHere are example questions (do not copy text, just mirror style):\n" + style_context[:8000]) if style_context else ""
    image_rules = (
        "- If an image would materially improve clarity (e.g., geometry diagrams, coordinate plots), set needsImage=true and provide a concise image_prompt describing a simple, high-contrast, black-and-white diagram. Otherwise set needsImage=false and image_prompt=null.\n"
        if allow_images else
        "- Do not include any images. Set needsImage=false and image_prompt=null.\n"
    )
    return (
        f"Task: Create {count} original MCQs.\n"
        f"Subject: {triple.subject}\nUnit: {triple.unit}\nTopic: {triple.topic}\n"
        "Rules:\n"
        "- Keep questions solvable without images.\n"
        + image_rules +
        "- Each question must be squarely within the specified topic.\n"
        "- Do NOT quote, paraphrase, or repeat any example content verbatim; produce new questions.\n"
        "- Preserve any LaTeX you output (e.g., $...$, \\(...\\), \\[...\\]).\n"
        "- Exactly 4 options; exactly one isCorrect=true.\n"
        "- Difficulty in {easy,moderate,hard}.\n"
        "- Explanations must be correct and concise.\n"
        "Return ONLY a JSON array of items with keys: question, instruction, difficulty, options[{text,isCorrect}], explanation, needsImage:boolean, image_prompt:string|null.\n"
        f"Produce exactly {count} items.\n" + ctx + "\nConstraints:\n- Return only the JSON array and nothing else.\n- Do not include the example text in your output.\n"
    )


def parse_items(text: str, expected_count: int) -> List[GeneratedItem]:
    text = (text or "").strip()
    if not text:
        raise ValueError("Model returned empty content (check model access/quota)")
    # Remove common code fences
    if text.startswith("```"):
        # strip first fence
        first_nl = text.find("\n")
        if first_nl != -1:
            text = text[first_nl+1:]
        if text.endswith("```"):
            text = text[:-3]
        text = text.strip()
    # Extract bracketed JSON array if the model adds prose
    start = text.find('[')
    end = text.rfind(']')
    if start != -1 and end != -1 and end > start:
        text = text[start:end+1]
    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        fixed = _fix_invalid_json_escapes(text)
        data = json.loads(fixed)
    if not isinstance(data, list):
        raise ValueError("Model did not return a JSON array")
    sanitized: List[GeneratedItem] = []
    for raw in data[:expected_count]:
        if not isinstance(raw, dict):
            raw = {}
        q = str(raw.get("question", "") or "").strip()
        instr = str(raw.get("instruction", "Choose the best answer.") or "Choose the best answer.").strip()
        diff_raw = str(raw.get("difficulty", "moderate") or "moderate").lower()
        diff = diff_raw if diff_raw in {"easy", "moderate", "hard"} else "moderate"
        exp = str(raw.get("explanation", "") or "").strip()
        needs_img = bool(raw.get("needsImage", False))
        img_prompt = raw.get("image_prompt", None)
        img_prompt_str = (str(img_prompt).strip() if isinstance(img_prompt, str) else "")
        options_raw = raw.get("options", [])
        opts: List[Option] = []
        if isinstance(options_raw, list):
            for o in options_raw:
                if isinstance(o, dict):
                    text_val = str(o.get("text", "") or "").strip()
                    is_corr = bool(o.get("isCorrect", False))
                    opts.append(Option(text=text_val, isCorrect=is_corr))
                else:
                    opts.append(Option(text=str(o), isCorrect=False))
        # pad to 4
        while len(opts) < 4:
            opts.append(Option(text=f"Option {len(opts)+1}", isCorrect=False))
        # trim to 4
        opts = opts[:4]
        # ensure exactly one correct
        num_correct = sum(1 for o in opts if o.isCorrect)
        if num_correct == 0:
            opts[0].isCorrect = True
        elif num_correct > 1:
            seen = False
            for o in opts:
                if o.isCorrect and not seen:
                    seen = True
                else:
                    o.isCorrect = False
        sanitized.append(GeneratedItem(
            question=q,
            instruction=instr,
            difficulty=diff,
            options=opts,
            explanation=exp,
            needsImage=needs_img if img_prompt_str else False,
            image_prompt=img_prompt_str if needs_img else "",
        ))
    if len(sanitized) != expected_count:
        raise ValueError("Insufficient valid items")
    return sanitized


def render_block(order_idx: int, triple: TopicTriple, item: GeneratedItem) -> str:
    lines: List[str] = []
    lines.append(f"@question {_clean_field(item.question)}")
    lines.append(f"@instruction {_clean_field(item.instruction)}")
    lines.append(f"@difficulty {_clean_field(item.difficulty)}")
    lines.append(f"@Order {order_idx}")

    marked = False
    for opt in item.options[:4]:
        if opt.isCorrect and not marked:
            lines.append(f"@@option {_clean_field(opt.text)}")
            marked = True
        else:
            lines.append(f"@option {_clean_field(opt.text)}")

    if item.needsImage and item.image_prompt:
        lines.append(f"@image_prompt {_clean_field(item.image_prompt)}")
    lines.append("@explanation")
    lines.append(_clean_field(item.explanation))
    lines.append(f"@subject {_clean_field(triple.subject)}")
    lines.append(f"@unit {_clean_field(triple.unit)}")
    lines.append(f"@topic {_clean_field(triple.topic)}")
    lines.append(f"@plusmarks 1")
    return "\n".join(lines) + "\n\n"


def to_docx_bytes(blocks_text: str, title: str, description: str) -> bytes:
    from docx import Document
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    if description:
        doc.add_paragraph(description)
    for para in blocks_text.strip().split("\n\n"):
        doc.add_paragraph(para)
        doc.add_paragraph("")
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def main():
    # Load .env first so env vars are available
    load_dotenv()
    st.set_page_config(page_title="Gemini MCQ Generator", layout="wide")
    st.title("Gemini MCQ Generator")

    with st.sidebar:
        # Prefer values from .env/.environment, but allow override via UI
        default_api_key = read_first_env(["GOOGLE_API_KEY", "GEMINI_API_KEY", "OPENAI_API_KEY", "API_KEY", "SPILEY", "spiley"], "")
        default_model = read_first_env(["GEMINI_MODEL", "MODEL", "OPENAI_MODEL", "model"], "gemini-2.5-flash")
        default_base = read_first_env(["GEMINI_BASE_URL", "BASE_URL", "OPENAI_BASE_URL", "base_url"], DEFAULT_BASE_URL)

        api_key = st.text_input("API Key", type="password", value=default_api_key)
        model = st.text_input("Model", value=default_model)
        base_url = st.text_input("Base URL", value=default_base)
        temperature = st.slider("Temperature", 0.0, 1.2, 0.6, 0.05)
        reasoning_effort = st.selectbox("Thinking mode (reasoning)", ["none", "low", "medium", "high"], index=0)
        attach_examples = st.checkbox("Attach base .md as style context (may echo)", value=False)
        reject_copies = st.checkbox("Reject questions that appear copied from base file", value=True)
        allow_images = st.checkbox("Allow images when helpful (adds @image_prompt)", value=False)
        count = st.number_input("Total questions", min_value=2, max_value=200, value=25, step=1)
        seed = st.number_input("Seed (reproducible)", min_value=0, max_value=2**31-1, value=123, step=1)
        title = st.text_input("Assessment title", value="Quantitative Math Assessment")
        description = st.text_area("Assessment description", value="Auto-generated MCQs across the curriculum")

    st.subheader("Example questions (.md, .txt)")
    uploaded = st.file_uploader("Upload your Example Questions File (.md, .txt)", type=["md", "txt"], accept_multiple_files=False)
    base_text = ""
    if uploaded is not None:
        base_text = uploaded.read().decode("utf-8", errors="ignore")
    else:
        # fallback: use existing file if present
        default_md = os.path.join(os.getcwd(), "ML-Official-T1-Math-01.md")
        if os.path.exists(default_md):
            with open(default_md, "r", encoding="utf-8") as f:
                base_text = f.read()

    if st.button("Generate"):
        if not api_key:
            st.error("Please provide GOOGLE_API_KEY")
            st.stop()
        triples = load_curriculum()
        plan = plan_distribution(triples, int(count), int(seed))
        client = build_client(base_url, api_key)

        system_prompt = build_system_prompt()
        all_blocks: List[str] = []
        preview_entries: List[Dict[str, Any]] = []
        order_idx = 1

        for p in plan:
            triple: TopicTriple = p["triple"]
            c = p["count"]
            if c <= 0:
                continue
            style_ctx = base_text if attach_examples else ""
            user_prompt = build_user_prompt(triple, c, style_ctx, allow_images)
            try:
                raw = call_model(client, model, system_prompt, user_prompt, float(temperature), reasoning_effort)
                items = parse_items(raw, c)
                # Optional duplicate detection against base file
                if reject_copies and base_text:
                    base_norm = _normalize_for_match(base_text)
                    dups = []
                    for idx, it in enumerate(items):
                        if _normalize_for_match(it.question) and _normalize_for_match(it.question) in base_norm:
                            dups.append(idx + 1)
                    if dups:
                        raise ValueError(f"Detected {len(dups)} question(s) possibly copied from base file at positions: {dups}")
            except Exception as e:
                st.error(f"Generation failed for {triple.subject} / {triple.unit} / {triple.topic}.\nModel: {model}\nBase URL: {base_url}\nReasoning: {reasoning_effort}\nParse hint: Ensure the model returns a pure JSON array as instructed.\nError: {e}")
                st.code(user_prompt)
                st.stop()

            for it in items:
                all_blocks.append(render_block(order_idx, triple, it))
                preview_entries.append({
                    "order": order_idx,
                    "subject": triple.subject,
                    "unit": triple.unit,
                    "topic": triple.topic,
                    "question": it.question,
                    "instruction": it.instruction,
                    "difficulty": it.difficulty,
                    "options": [{"text": o.text, "correct": o.isCorrect} for o in it.options[:4]],
                    "explanation": it.explanation,
                    "image_prompt": it.image_prompt if getattr(it, "needsImage", False) and getattr(it, "image_prompt", "") else "",
                })
                order_idx += 1

        header = []
        if title:
            header.append(f"@title {title}")
        if description:
            header.append(f"@description {description}")
        header_text = "\n".join(header) + ("\n\n" if header else "")

        full_text = header_text + "".join(all_blocks)
        st.success(f"Generated {count} questions.")
        st.subheader("Preview")
        with st.expander("Raw tag output", expanded=False):
            st.code(full_text, language="text")
        with st.expander("Structured preview", expanded=True):
            for entry in preview_entries:
                title_line = f"Q{entry['order']} [{entry['difficulty']}] — {entry['subject']} / {entry['unit']} / {entry['topic']}"
                st.markdown(f"**{title_line}**")
                st.markdown(f"Question: {_fix_inline_latex(entry['question'])}")
                st.markdown(f"Instruction: {entry['instruction']}")
                st.markdown("Options:")
                for opt in entry["options"]:
                    text = _fix_inline_latex(opt.get("text", ""))
                    suffix = " (correct)" if opt["correct"] else ""
                    st.markdown(f"- {text}{suffix}")
                if entry.get("image_prompt"):
                    st.markdown(f"Image prompt: {entry['image_prompt']}")
                with st.expander("Explanation", expanded=False):
                    st.write(entry["explanation"]) 
        st.download_button("Download .txt", data=full_text.encode("utf-8"), file_name="assessment.txt", mime="text/plain")

        # Optional DOCX export (in-memory only; not saved to disk)
        try:
            docx_bytes = to_docx_bytes(full_text, title, description)
            st.download_button(
                "Download .docx",
                data=docx_bytes,
                file_name="assessment.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception as e:
            st.info(f"DOCX export skipped: {e}")


if __name__ == "__main__":
    main()


